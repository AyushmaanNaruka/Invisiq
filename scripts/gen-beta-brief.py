# -*- coding: utf-8 -*-
"""
Generates: InvisiQ — Beta Launch Technical Brief (.docx)
Audience: (1) website developer, (2) app-launch / promo-video maker.
All content is grounded in the actual codebase + Beta Launch Plan. No invented facts.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette (authoritative, from globals.css / tailwind.config.ts) ──
NAVY      = "0B0E14"   # bg-overlay (deep navy)
NAVY_HDR  = "12161E"   # bg-header
TEAL      = "14B8A6"   # accent-primary
TEAL_DK   = "0D9488"
BLUE      = "3B82F6"
PURPLE    = "8B5CF6"
AMBER     = "FBBF24"
INK       = "1A2230"   # body text on white
SUBTLE    = "5B6B7F"   # secondary text
HAIR      = "D5DCE4"   # hairline
TINT_TEAL = "E7F7F4"
TINT_WARN = "FEF3C7"
TINT_DANG = "FDE7E7"
TINT_INFO = "EEF3FA"
WHITE     = "FFFFFF"

BODY_FONT = "Segoe UI"   # reliable on Windows; brand font is Inter (noted in doc)
MONO_FONT = "Consolas"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

def _shade(el, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    el.append(shd)

def cell_bg(cell, hexcolor):
    _shade(cell._tc.get_or_add_tcPr(), hexcolor)

def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)

def no_borders(table):
    tbl = table._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none"); e.set(qn("w:sz"), "0")
        borders.append(e)
    tbl.tblPr.append(borders)

def thin_borders(table, color=HAIR):
    tbl = table._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tbl.tblPr.append(borders)

def run(p, text, *, bold=False, italic=False, color=INK, size=10.5, font=BODY_FONT):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = font; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return r

def para(text=None, *, bold=False, italic=False, color=INK, size=10.5,
         space_before=0, space_after=6, align=None, font=BODY_FONT):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run(p, text, bold=bold, italic=italic, color=color, size=size, font=font)
    return p

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(2)
    run(p, text, bold=True, color=NAVY, size=17)
    # teal underline rule
    rule = doc.add_paragraph(); rule.paragraph_format.space_after = Pt(8)
    pPr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"18")
    bottom.set(qn("w:space"),"1"); bottom.set(qn("w:color"),TEAL)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
    run(p, text, bold=True, color=TEAL_DK, size=13)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    run(p, text, bold=True, color=NAVY, size=11)
    return p

def bullet(text, *, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.28 + 0.22*level)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run(p, bold_prefix, bold=True, color=NAVY)
        run(p, text)
    else:
        run(p, text)
    return p

def numbered(text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.30)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run(p, bold_prefix, bold=True, color=NAVY)
        run(p, text)
    else:
        run(p, text)
    return p

def callout(title, body_lines, *, tint=TINT_INFO, bar=BLUE, icon=""):
    """Single-cell shaded box with a colored left bar effect via title color."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    cell = t.cell(0,0)
    cell_bg(cell, tint)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    # left accent border
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),"single"); left.set(qn("w:sz"),"24")
    left.set(qn("w:space"),"0"); left.set(qn("w:color"),bar)
    borders.append(left); tcPr.append(borders)
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    run(p, f"{icon}{title}", bold=True, color=bar, size=10.5)
    for line in body_lines:
        bp = cell.add_paragraph(); bp.paragraph_format.space_after = Pt(2)
        run(bp, line, color=INK, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def table(headers, rows, *, widths=None, header_bg=NAVY, zebra=TINT_INFO, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    thin_borders(t)
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        cell_bg(hdr[i], header_bg)
        set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        run(p, htext, bold=True, color=WHITE, size=font_size)
    for ri, rowvals in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(rowvals):
            set_cell_margins(cells[ci])
            if zebra and ri % 2 == 1:
                cell_bg(cells[ci], zebra)
            p = cells[ci].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            # allow (text, color) tuples
            if isinstance(val, tuple):
                run(p, val[0], color=val[1], size=font_size, bold=val[2] if len(val)>2 else False)
            else:
                run(p, str(val), color=INK, size=font_size)
    if widths:
        for ci, w in enumerate(widths):
            for r in t.rows:
                r.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def code_box(lines):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0,0); cell_bg(cell, NAVY)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    cell.paragraphs[0].text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run(p, line, color="9FE9DD", size=9.5, font=MONO_FONT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def spacer(pts=4):
    doc.add_paragraph().paragraph_format.space_after = Pt(pts)

def pagebreak():
    doc.add_page_break()

# Wider margins -> more usable width
for s in doc.sections:
    s.left_margin = Inches(0.85); s.right_margin = Inches(0.85)
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)

# ════════════════════════════════════════════════════════════════════
#  COVER
# ════════════════════════════════════════════════════════════════════
cover = doc.add_table(rows=1, cols=1)
ccell = cover.cell(0,0)
cell_bg(ccell, NAVY)
set_cell_margins(ccell, top=520, bottom=520, left=320, right=320)
ccell.paragraphs[0].text = ""
p = ccell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run(p, "InvisiQ", bold=True, color=TEAL, size=40)
p2 = ccell.add_paragraph(); run(p2, "Your AI copilot that sees everything, but is seen by no one.",
                                italic=True, color="C7D2DE", size=13)
p3 = ccell.add_paragraph(); p3.paragraph_format.space_before = Pt(18)
run(p3, "BETA LAUNCH — TECHNICAL & BRAND BRIEF", bold=True, color=WHITE, size=15)
p4 = ccell.add_paragraph()
run(p4, "Reference pack for the website developer and the launch / promo-video producer.",
    color="9FB0C2", size=10.5)
p5 = ccell.add_paragraph(); p5.paragraph_format.space_before = Pt(14)
run(p5, "Product: InvisiQ  ·  Platform: Windows 10 (2004+) / 11  ·  Beta build v1.2.0", color="C7D2DE", size=9.5)
p6 = ccell.add_paragraph()
run(p6, "Owner: Ayushmaan Singh Naruka  ·  Prepared: 8 June 2026  ·  Status: Pre-launch (beta)", color="9FB0C2", size=9.5)

spacer(6)
callout(
    "How to read this document",
    [
        "Sections 1–6 = product truth (what InvisiQ is, what it does, who it's for). Read first.",
        "Sections 7–10 = the beta mechanics, supported models, shortcuts, system requirements.",
        "Sections 11–13 = PRIVACY/LEGAL, BRAND & VISUAL IDENTITY — most relevant to the website developer.",
        "Section 14 = PROMO-VIDEO BRIEF (storyboard, captions, assets) — for the video producer.",
        "Section 15 = WEBSITE STRUCTURE & messaging guardrails. Section 17 lists items the owner must still provide.",
        "Every claim here is taken from the actual InvisiQ build. Where something is undecided it is marked “TBD — owner to provide.” Do not invent specs.",
    ],
    tint=TINT_TEAL, bar=TEAL_DK, icon="✓  "
)

pagebreak()

# ════════════════════════════════════════════════════════════════════
#  1. PRODUCT AT A GLANCE
# ════════════════════════════════════════════════════════════════════
h1("1. Product at a glance")
table(
    ["Field", "Value"],
    [
        ("Product name", ("InvisiQ", INK, True)),
        ("Tagline", "Your AI copilot that sees everything, but is seen by no one."),
        ("One-liner", "An invisible AI desktop overlay — invisible to Zoom, Teams, Meet, OBS and proctoring tools, yet always on top for you."),
        ("Category", "Invisible AI desktop assistant / screen-aware copilot (Windows overlay app)"),
        ("Platform (beta)", "Windows 10 version 2004+ and Windows 11 only. No macOS / Linux in beta."),
        ("How AI is provided", "Bring-Your-Own-Key (BYOK): the user connects their own OpenAI / Anthropic / Google API key. Cloud only."),
        ("Price (beta)", "Free 14-day trial after Google sign-in. The user pays only their own AI provider. Paid InvisiQ plans come later."),
        ("Distribution", "Single portable .exe — no installer, no admin rights. (Beta is unsigned; see §10.)"),
        ("Owner", "Ayushmaan Singh Naruka"),
    ],
    widths=[1.7, 5.0]
)

h2("Elevator pitch (use this as the website hero copy starting point)")
para("InvisiQ is a desktop overlay that floats on top of everything you do — and is completely invisible "
     "to screen sharing, screen recording, screenshots, and meeting/monitoring software. You can read a question "
     "on screen, ask an AI about it, and get a streaming answer in a sleek panel that no one else on the call (or "
     "in the recording) can see. It is controlled entirely by keyboard shortcuts and connects to the AI providers "
     "you already use. Nothing about your screen content ever leaves your machine except the questions you choose to ask.")

h2("What it is / what it is NOT")
two = doc.add_table(rows=1, cols=2); thin_borders(two)
c1, c2 = two.rows[0].cells
cell_bg(c1, TINT_TEAL); cell_bg(c2, TINT_DANG)
set_cell_margins(c1); set_cell_margins(c2)
c1.paragraphs[0].text=""; c2.paragraphs[0].text=""
run(c1.paragraphs[0], "IS", bold=True, color=TEAL_DK, size=11)
for line in [
    "A privacy-first, screen-aware AI overlay for Windows.",
    "Invisible to capture APIs (Snipping Tool, OBS, Zoom/Teams/Meet share, recordings).",
    "BYOK — you use your own AI keys; direct calls to the official AI providers.",
    "Local-first: your settings, chats and keys stay encrypted on your machine.",
    "Keyboard-driven, always-on-top, multi-monitor aware.",
]:
    bp=c1.add_paragraph(); bp.paragraph_format.space_after=Pt(2); run(bp,"• "+line,size=9.5)
run(c2.paragraphs[0], "IS NOT", bold=True, color="B23B3B", size=11)
for line in [
    "Not a cloud service that stores your screens — screenshots never leave the device.",
    "Not a local-LLM tool — Ollama/offline models are removed in the beta (cloud only).",
    "Not a macOS/Linux app yet — Windows only for the beta.",
    "Not free of an AI bill — the user still pays their own provider for tokens.",
    "Not a guarantee against a physical camera pointed at the screen (see §6 caveats).",
]:
    bp=c2.add_paragraph(); bp.paragraph_format.space_after=Pt(2); run(bp,"• "+line,size=9.5)
spacer(4)

# ════════════════════════════════════════════════════════════════════
#  2. WHAT INVISIQ DOES
# ════════════════════════════════════════════════════════════════════
h1("2. What InvisiQ does — core capabilities")
para("Every capability below is implemented and shipping in the beta build. Group A is the headline story for the "
     "website and video. Group B is the supporting feature depth.", italic=True, color=SUBTLE)

h2("Group A — the headline capabilities")
table(
    ["Capability", "What the user experiences"],
    [
        ("Invisible overlay", "A floating panel that is excluded from all screen capture at the OS level — it shows on your physical monitor but is stripped out of any shared screen, recording, or screenshot."),
        ("Ask about your screen", "Capture the full screen, a dragged region, or an in-overlay snip; the image goes to an AI vision model and the answer streams back in the panel."),
        ("Streaming AI chat", "Token-by-token answers with Markdown, syntax-highlighted code blocks, and one-click copy. Multi-turn context is preserved."),
        ("Stealth typing", "Type into the overlay from any app without it stealing focus — keystrokes are captured invisibly so the panel never becomes the foreground window a monitor would notice."),
        ("Smart modes", "One-tap personas: General, Coding, Meeting, Solve — plus user-defined custom modes. Each tunes the AI's behavior for the task."),
        ("Works under monitoring", "Invisible to Zoom/Teams/Meet screen-share, OBS, Discord, Snipping Tool, and to foreground-window monitoring used by proctoring/secure-browser tools."),
        ("Panic & instant hide", "Escape hides the overlay instantly; a panic hotkey exits capture and hides everything. Nothing lingers on screen."),
    ],
    widths=[1.7, 5.0]
)

h2("Group B — supporting features (built, available in beta)")
b = [
    ("Conversation history", "Auto-saved chats with search and export to JSON / Markdown / TXT / PDF."),
    ("Memory (RAG)", "A built-in TF-IDF memory auto-extracts facts and feeds relevant context back into prompts."),
    ("Smart paste", "Drop an AI answer or code block straight into the active app (overlay hides, pastes, restores)."),
    ("Voice + meeting assist", "Speech-to-text (Web Speech / Whisper), live transcript, system-audio capture, auto-question detection and suggestions."),
    ("Prompt templates", "20+ built-in templates across 8 categories, plus your own."),
    ("Companion mode", "Pair a phone over a local QR/WebSocket link to drive the overlay from a second device."),
    ("Multi-monitor", "Hot-plug detection; screenshots and the region selector target the right display."),
    ("Themes & cost meter", "Dark and light themes; per-request / per-session token & cost tracking in the status bar."),
    ("Auto-update", "Self-updates via the release feed so beta users stay current."),
]
for name, desc in b:
    bullet(desc, bold_prefix=name + " — ")
spacer(2)

# ════════════════════════════════════════════════════════════════════
#  3. HOW IT WORKS (the invisibility)
# ════════════════════════════════════════════════════════════════════
h1("3. How the invisibility works (plain + technical)")
para("This is the one mechanism the whole product rests on; describe it accurately so the marketing is truthful.")

h2("Plain-language version (safe for website/video)")
para("Windows draws every window through a compositor (the Desktop Window Manager). InvisiQ asks Windows to put its "
     "overlay on a special layer that is sent to your monitor but withheld from anything that captures the screen. "
     "So the pixels reach your eyes, but not the screen-share, the recorder, or the screenshot. It is a documented "
     "Windows privacy feature — not a hack or a driver exploit.")

h2("Technical version (for credibility / FAQ)")
bullet("setContentProtection(true) in Electron calls the Windows API SetWindowDisplayAffinity(hwnd, WDA_EXCLUDEFROMCAPTURE) — introduced in Windows 10 v2004 for privacy-sensitive apps.", bold_prefix="Capture exclusion: ")
bullet("A second stealth layer makes the overlay non-activatable (WS_EX_NOACTIVATE) so a single click can’t make it the foreground window — this is what defeats foreground-window monitoring used by secure browsers / proctoring tools, not just visual capture.", bold_prefix="Focus stealth: ")
bullet("Because the window won’t take focus, free-form typing is delivered through an out-of-process keystroke capture pipeline (“stealth typing”) instead of normal window focus.", bold_prefix="Stealth typing: ")
bullet("A watchdog re-applies content protection every 2 seconds in case the state ever drops.", bold_prefix="Self-healing: ")
bullet("The process presents itself with a low-profile system-style name in Task Manager so it doesn’t stand out in a process list.", bold_prefix="Process camouflage: ")

callout("Honest limits (keep marketing truthful)",
    [
        "A physical camera pointed at the monitor will still see the overlay — capture exclusion is a software pipeline, not magic.",
        "macOS 15+ weakened the equivalent flag; the beta is Windows-only, so this is out of scope for now.",
        "Kernel-level enterprise DLP/endpoint agents can enumerate running processes; the disguise reduces but cannot guarantee non-detection on a fully managed device.",
        "Do NOT claim “undetectable by anything” or “100% guaranteed invisible everywhere.” The accurate claim is: invisible to screen capture / sharing / recording APIs and to foreground-window monitoring on Windows 10 2004+.",
    ],
    tint=TINT_WARN, bar=AMBER, icon="⚠  ")

pagebreak()

# ════════════════════════════════════════════════════════════════════
#  4. TARGET USERS & USE CASES
# ════════════════════════════════════════════════════════════════════
h1("4. Target users & use cases")
para("These use cases are real and supported by the product. Section 15 covers HOW to frame them publicly without "
     "creating legal/brand risk — read both together before writing site copy or a script.")
table(
    ["Use case", "What InvisiQ does in it"],
    [
        ("Live video meetings", "Privately summarize a shared document/slide, get talking points, or look something up — invisible to everyone on the call, even if you’re the one sharing."),
        ("Interview & assessment practice", "Rehearse coding/aptitude problems with instant approach + Big-O + worked solutions in Coding/Solve mode."),
        ("Learning & study companion", "Capture a diagram, equation, or passage and get a step-by-step explanation; save the thread for revision."),
        ("Real-time research", "Pull quick context, definitions, or data analysis without alt-tabbing away from what you’re doing."),
        ("Sales / client calls", "Generate talking points and objection handling live, without it showing on a shared screen."),
        ("Accessibility / focus", "A keyboard-first, always-available assistant that doesn’t clutter the desktop or taskbar."),
    ],
    widths=[1.9, 4.8]
)

# ════════════════════════════════════════════════════════════════════
#  5. THE BETA
# ════════════════════════════════════════════════════════════════════
h1("5. The beta — what a user actually gets")
table(
    ["Aspect", "Beta behavior"],
    [
        ("Sign-in", "Google sign-in is required (opens the system browser; nothing to type into the app). Stays signed in on later launches."),
        ("Trial", "A 14-day free trial begins on first authenticated launch. The countdown shows in a banner inside the app."),
        ("After 14 days", "AI features hard-lock. The app shows “Your trial has ended — paid plans are coming soon.” The user can still sign out."),
        ("Anti-cheat on the trial", "The 14 days are enforced by InvisiQ’s server clock. Changing the system clock, reinstalling, clearing data, or making a new device install does not reset the trial."),
        ("AI keys", "BYOK — the user pastes their own OpenAI / Anthropic / Google key. Keys are encrypted on the device and used to call the providers directly."),
        ("Offline", "Needs an internet check to confirm the trial. If it can’t verify, it shows “Can’t verify your trial — reconnect and try again.” (fail-closed)."),
        ("Cost to user", "InvisiQ is free during the beta; the only cost is the user’s own AI provider usage."),
    ],
    widths=[1.7, 5.0]
)

h2("Trial / launch states the video & site should depict accurately")
bullet("Sign-in screen — “Continue with Google.”", bold_prefix="1. ")
bullet("Trial verified — the app opens with a small “N days left” banner.", bold_prefix="2. ")
bullet("Beta data notice (T&C gate) — a one-time “Before you start” screen the user must accept (see §11).", bold_prefix="3. ")
bullet("Onboarding — a 3-step wizard: add an API key, see the hotkeys, run a stealth self-test.", bold_prefix="4. ")
bullet("In use — the overlay with streaming answers; stealth typing ON shows a glowing input ring.", bold_prefix="5. ")
bullet("Trial ended — a lock screen with “paid plans coming soon.”", bold_prefix="6. ")

# ════════════════════════════════════════════════════════════════════
#  6. FIRST-RUN FLOW (storyboard-ready)
# ════════════════════════════════════════════════════════════════════
h1("6. First-run user journey (exact order)")
para("This is the real screen order in the app (useful for the onboarding section of the site and for the video).")
code_box([
    "Launch  ->  Google sign-in  ->  Trial check (server, 14 days)",
    "        ->  “Before you start” beta-data notice (accept once)",
    "        ->  3-step onboarding (API key / hotkeys / stealth test)",
    "        ->  Overlay ready  (Ctrl+Shift+G toggles it anytime)",
    "",
    "After 14 days  ->  Lock screen (“Your trial has ended”)",
])

pagebreak()

# ════════════════════════════════════════════════════════════════════
#  7. SUPPORTED PROVIDERS & MODELS
# ════════════════════════════════════════════════════════════════════
h1("7. Supported AI providers & models (beta)")
callout("Accuracy note", [
    "The public GitHub README still lists “Ollama / local AI.” That is pre-beta. The beta build is CLOUD-ONLY and "
    "Ollama/local models have been removed. Do not advertise local/offline AI for the beta.",
], tint=TINT_WARN, bar=AMBER, icon="⚠  ")
table(
    ["Provider", "Model", "Vision", "Context", "Speed"],
    [
        ("OpenAI", "GPT-4o", "Yes", "128K", "Medium"),
        ("OpenAI", "GPT-4o Mini", "Yes", "128K", "Fast"),
        ("OpenAI", "o3-mini (reasoning)", "No", "200K", "Slow"),
        ("Anthropic", "Claude Sonnet 4", "Yes", "200K", "Medium"),
        ("Anthropic", "Claude Haiku 4.5", "Yes", "200K", "Fast"),
        ("Google", "Gemini 2.0 Flash", "Yes", "1M", "Fast"),
        ("Google", "Gemini 2.5 Pro", "Yes", "1M", "Medium"),
    ],
    widths=[1.2, 2.3, 0.9, 1.0, 1.0]
)
para("Users switch model and provider from a dropdown in the header, mid-conversation. “Vision” models read "
     "screenshots directly; o3-mini is text/reasoning only.", italic=True, color=SUBTLE, size=9.5)

# ════════════════════════════════════════════════════════════════════
#  8. KEYBOARD SHORTCUTS
# ════════════════════════════════════════════════════════════════════
h1("8. Keyboard shortcuts (default, all customizable)")
para("These are the defaults the beta build actually registers. (They are user-editable in Settings → Hotkeys.)",
     italic=True, color=SUBTLE, size=9.5)
table(
    ["Shortcut", "Action"],
    [
        ("Ctrl + Shift + G", "Toggle overlay visibility"),
        ("Ctrl + Alt + S", "Capture full screen → send to AI"),
        ("Ctrl + Alt + R", "Capture a dragged region → send to AI"),
        ("Ctrl + Alt + A", "Focus the text input"),
        ("Ctrl + Alt + C", "Copy the last AI response"),
        ("Ctrl + Shift + V", "Paste the last AI response into the active app"),
        ("Ctrl + Alt + N", "New conversation"),
        ("Ctrl + Shift + I", "Toggle stealth typing / capture mode"),
        ("Ctrl + Shift + P", "Toggle click-through (mouse passes through the overlay)"),
        ("Ctrl + Alt + ]  /  [", "Next / previous model"),
        ("Ctrl + Shift + Q", "Panic — exit capture, uninstall the keyboard hook, hide overlay"),
        ("Escape", "Hide the overlay immediately"),
    ],
    widths=[2.1, 4.6]
)

# ════════════════════════════════════════════════════════════════════
#  9. SYSTEM REQUIREMENTS
# ════════════════════════════════════════════════════════════════════
h1("9. System requirements")
table(
    ["Requirement", "Detail"],
    [
        ("Operating system", "Windows 10 version 2004 (build 19041) or later, or Windows 11. Required for the WDA_EXCLUDEFROMCAPTURE capture-exclusion API."),
        ("macOS / Linux", "Not supported in the beta."),
        ("Display", "1280×720 or higher; multi-monitor supported."),
        ("Network", "Internet connection required (trial verification + AI API calls)."),
        ("Install footprint", "Portable .exe — no installer, no admin rights, no registry entries, no Start-Menu shortcut."),
        ("AI account", "At least one API key from OpenAI, Anthropic, or Google AI Studio."),
        ("Performance targets", "Cold start under ~3s; idle RAM under ~150 MB (design targets, not guarantees)."),
    ],
    widths=[1.7, 5.0]
)

pagebreak()

# ════════════════════════════════════════════════════════════════════
#  10. DISTRIBUTION & THE SMARTSCREEN NOTE
# ════════════════════════════════════════════════════════════════════
h1("10. Distribution & the unsigned-beta note (website-critical)")
para("The beta ships as a single unsigned portable .exe. This is a deliberate beta decision (signing comes once a legal "
     "company entity exists). It has a concrete consequence the website MUST prepare users for:")
callout("SmartScreen “unknown publisher” warning",
    [
        "On first run, Windows SmartScreen may show “Windows protected your PC — unknown publisher.” This is expected for an unsigned beta and does not mean the app is unsafe.",
        "The download page should include a short “Why do I see a Windows warning?” note with the click path: More info → Run anyway.",
        "Do not tell users to disable Defender or add exclusions — just the More info → Run anyway step.",
        "TBD — owner to provide the actual download host/URL (currently distributed via the release feed; a public download page is part of launch).",
    ],
    tint=TINT_WARN, bar=AMBER, icon="⚠  ")
para("Auto-update is built in, so once a user is on a build it keeps itself current.", color=SUBTLE, size=9.5, italic=True)

# ════════════════════════════════════════════════════════════════════
#  11. PRIVACY, DATA & LEGAL  (website-critical)
# ════════════════════════════════════════════════════════════════════
h1("11. Privacy, data & legal — must be reflected on the site")
callout("This section is non-negotiable for the website developer",
    ["The site needs a Privacy Policy and Beta Terms before public launch, and they must match exactly what the app does below. "
     "Getting this wrong is the single biggest reputational/legal risk for a product whose pitch is “seen by no one.”"],
    tint=TINT_DANG, bar="B23B3B", icon="⚠  ")

h2("What the app stores during the beta (accurate)")
bullet("The TEXT of the prompts the user sends is logged to InvisiQ’s backend, to learn what to build. This is disclosed in-app and the user must accept it before use.", bold_prefix="Prompt text: ")
bullet("Screenshots and screen contents are NEVER stored or uploaded. Image attachments are recorded only as a yes/no flag; the image itself is not kept.", bold_prefix="No screen content: ")
bullet("API keys and obvious personal info (emails, phone numbers, card-like numbers) are stripped server-side before a prompt is stored.", bold_prefix="Redaction: ")
bullet("Stored prompt text is deleted after 30 days. The user can wipe their own data anytime via Settings → Privacy (“Delete my data”).", bold_prefix="Retention: ")
bullet("Privacy-safe product analytics (e.g. app launched, signed in, message sent with a length bucket, trial expired) are also collected.", bold_prefix="Analytics: ")
bullet("API keys, chats and settings live encrypted on the user’s machine; AI calls go directly to the official providers.", bold_prefix="Local-first: ")

h2("The exact in-app disclosure the user accepts (“Before you start” gate)")
para("Reproduce this verbatim wording on the website’s privacy page so they match. This is the real screen copy:",
     italic=True, color=SUBTLE, size=9.5)
callout("“Before you start” (current in-app text)",
    [
        "During the beta, InvisiQ stores the text of the prompts you send, so we can understand what to build and improve the product.",
        "We store text only — never your screenshots or screen contents. API keys and obvious personal info are stripped before storage. Prompt data is deleted after 30 days, and you can wipe yours anytime in Settings → Privacy.",
    ],
    tint=TINT_INFO, bar=BLUE, icon="")

callout("Open legal items — owner to resolve before public launch",
    [
        "The in-app T&C copy is a PLACEHOLDER. Counsel-reviewed Beta Terms + Privacy Policy are required before a public download.",
        "Because prompt-text collection is mandatory (bundled into the terms), GDPR (EU) and India’s DPDP Act 2023 are a real consideration — counsel may advise geo-limiting the beta. Owner decision.",
        "The Privacy Policy must state: what is stored (prompt text + analytics), what is never stored (screens), retention (30 days), how to delete, and who can access it.",
    ],
    tint=TINT_WARN, bar=AMBER, icon="⚠  ")

pagebreak()

# ════════════════════════════════════════════════════════════════════
#  12. BRAND & VISUAL IDENTITY  (website + video)
# ════════════════════════════════════════════════════════════════════
h1("12. Brand & visual identity")
para("These tokens are pulled directly from the live app theme (globals.css / tailwind.config.ts). Match them so the "
     "website and video feel like the product. Note: some older internal docs list a different palette "
     "(e.g. #1a1a2e / #00B894) — those are SUPERSEDED. The values below are authoritative.", italic=True, color=SUBTLE, size=9.5)

h2("Color — dark theme (the default / hero look)")
table(
    ["Role", "Hex", "RGB", "Use"],
    [
        ("Accent / primary (teal)", ("#14B8A6", TEAL_DK, True), "20 184 166", "Primary brand color: logo, CTAs, glows, active states"),
        ("Deep navy (surface)", "#0B0E14", "11 14 20", "Overlay/page background — the signature dark canvas"),
        ("Header surface", "#12161E", "18 22 30", "Slightly lighter panel/header bars"),
        ("Input surface", "#161B24", "22 27 36", "Inputs, secondary panels"),
        ("Code background", "#090C12", "9 12 18", "Code blocks"),
        ("Text primary", "#E2E8F0", "226 232 240", "Headlines / body on dark"),
        ("Text secondary", "#64748B", "100 116 139", "Muted captions / labels"),
        ("Accent blue", "#3B82F6", "59 130 246", "Links / secondary accent"),
        ("Accent purple", "#8B5CF6", "139 92 246", "Coding-mode accent"),
        ("Accent amber", "#FBBF24", "251 191 36", "Solve-mode accent / warnings"),
        ("Success green", "#22C55E", "34 197 94", "“Stealth ON” shield, connected dot"),
        ("Error red", "#EF4444", "239 68 68", "Errors / “you are sharing” badge"),
        ("User bubble blue", "#2563EB", "37 99 235", "User message bubble"),
    ],
    widths=[1.9, 1.1, 1.3, 2.4], font_size=9
)

h2("Color — light theme (the app also ships a light mode)")
table(
    ["Role", "Hex"],
    [
        ("Accent / primary (teal, deeper)", "#0D9488"),
        ("Page background", "#F2F4F8"),
        ("Surface / cards", "#FFFFFF"),
        ("Text primary", "#0F172A"),
    ],
    widths=[2.6, 1.4], font_size=9.5
)

h2("Typography")
bullet("Inter — all UI / headings / body. (SF Pro Display / system-ui as fallback.)", bold_prefix="Primary: ")
bullet("JetBrains Mono — code blocks and anything mono. (Fira Code / Consolas fallback.)", bold_prefix="Mono: ")
bullet("Base UI size 13px; the app uses a compact scale (11–22px). Headlines on the site can go larger, but keep Inter.", bold_prefix="Sizing: ")

h2("Shape, depth & motion (so the site/video feel native to the app)")
bullet("12px corner radius on the main panel; 4–20px scale elsewhere.", bold_prefix="Radius: ")
bullet("Panel sits at ~85% opacity — a glassy, slightly translucent dark surface.", bold_prefix="Opacity: ")
bullet("Teal glow: box-shadow 0 0 20px rgba(20,184,166,0.20), 0 0 40px rgba(20,184,166,0.08).", bold_prefix="Signature glow: ")
bullet("Easing cubic-bezier(0.16, 1, 0.3, 1) (“out-expo”). Show animation 200ms ease-out (fade + scale 0.95→1.0); hide 150ms.", bold_prefix="Motion: ")
bullet("Always respect prefers-reduced-motion; keep motion calm/restrained, no harsh flashes or strobe.", bold_prefix="Restraint: ")

h2("Tone of voice")
para("Calm, premium, “stealth-tech.” Confident but not boastful; privacy-forward. Short, declarative lines. "
     "Avoid hype words like “undetectable,” “unhackable,” or “guaranteed.” Lead with control and privacy.")

h2("Signature visual motifs")
bullet("The teal “ghost” mark (the InvisiQ logo).", )
bullet("The glowing teal input ring = “stealth typing” is live.")
bullet("A green shield = stealth ON; a red “you are sharing your screen” badge for contrast.")
bullet("The “what the capture sees vs. what you see” framing (a dashed frame that the overlay sits outside of).")

pagebreak()

# ════════════════════════════════════════════════════════════════════
#  13. ASSETS THAT ALREADY EXIST
# ════════════════════════════════════════════════════════════════════
h1("13. Existing assets to hand to the designers")
para("The repo already contains pixel-faithful HTML mockups of the real UI plus a prepared prompt pack. Give these to "
     "both the website designer and the video producer — they are the visual source of truth.")
table(
    ["File (in /claude-design)", "Shows", "Use as"],
    [
        ("invisiq-overlay-empty.html", "Overlay just opened — logo, tagline, hotkey hints", "“App launches” beat / empty state"),
        ("invisiq-overlay.html", "Active chat: user Q + AI answer w/ code, stealth typing ON, proctor badge", "Core product UI"),
        ("invisiq-hero-floating.html", "Overlay floating over a screen-shared call; dashed frame = what the capture sees", "Establishing / payoff shot"),
        ("CLAUDE-DESIGN-PROMPTS.md", "Ready-to-use prompts for an AI motion/explainer tool", "Video generation brief"),
    ],
    widths=[2.4, 2.9, 1.4], font_size=9
)
para("These mockups use the exact teal #14B8A6 on deep navy #0B0E14, Inter UI / JetBrains Mono, 12px corners, 85% "
     "glassy panel. Tell designers to match them rather than redesign the interface.", italic=True, color=SUBTLE, size=9.5)

# ════════════════════════════════════════════════════════════════════
#  14. PROMO VIDEO BRIEF
# ════════════════════════════════════════════════════════════════════
h1("14. Promo / launch video brief")
para("A 25-second hero explainer is the recommended primary deliverable, with an 8-second silent loop for the website "
     "hero. Storyboard below is adapted from the prepared prompt pack and the real first-run flow.")

h2("Primary: 25-second hero explainer (16:9, 1920×1080, MP4 + web loop)")
table(
    ["Time", "On screen", "Caption"],
    [
        ("0–4s", "A normal video call / screen-share with a red “You are sharing your screen” badge. Calm, ordinary.", "(none)"),
        ("4–9s", "The InvisiQ overlay fades + scales in from bottom-right (200ms ease-out, 0.95→1.0).", "“Meet InvisiQ.”"),
        ("9–16s", "Active overlay: a question types into the glowing teal input; an AI answer with a code block streams in token-by-token; blinking caret.", "“Ask anything. Get answers — instantly.”"),
        ("16–21s", "Pull back to the call. A dashed red frame “What the screen-share sees” sweeps across — the overlay is OUTSIDE it.", "“Invisible to Zoom, Teams, Meet, OBS and monitoring tools.”"),
        ("21–25s", "End card: centered InvisiQ ghost mark + wordmark, tagline, soft teal glow pulse.", "“Your AI copilot that sees everything, but is seen by no one.”"),
    ],
    widths=[0.8, 4.0, 1.9], font_size=9
)

h2("Secondary clips")
bullet("8-second silent loop for the website hero: input ring pulses → short AI reply streams in → green shield glints → soft teal glow breathes. Perfect loop, no captions. (16:9 + 1:1).", bold_prefix="Hero loop: ")
bullet("10-second “Type without being seen” spotlight: zoom into the input row, keyboard icon lights teal, placeholder is replaced by text appearing character-by-character; proctor badge “Monitored app detected — you’re invisible” slides up.", bold_prefix="Feature spotlight: ")
bullet("Static set: 1200×630 OG/social card, 1080×1080 square, and a clean product shot on a soft gradient.", bold_prefix="Social/OG: ")

h2("Video do / don’t")
two2 = doc.add_table(rows=1, cols=2); thin_borders(two2)
d1, d2 = two2.rows[0].cells
cell_bg(d1, TINT_TEAL); cell_bg(d2, TINT_DANG)
set_cell_margins(d1); set_cell_margins(d2)
d1.paragraphs[0].text=""; d2.paragraphs[0].text=""
run(d1.paragraphs[0], "DO", bold=True, color=TEAL_DK, size=11)
for line in [
    "Let the product motion (streaming text, glow, fade-in) carry it.",
    "Use real UI from the mockups; keep captions short.",
    "Show meetings, learning, interview prep, research as the contexts.",
    "Keep it calm/premium; smooth easing; reduced-motion friendly.",
]:
    bp=d1.add_paragraph(); bp.paragraph_format.space_after=Pt(2); run(bp,"• "+line,size=9.5)
run(d2.paragraphs[0], "DON’T", bold=True, color="B23B3B", size=11)
for line in [
    "Don’t frame it as a tool to cheat on a proctored exam (see §15).",
    "Don’t use the words “undetectable,” “unhackable,” “guaranteed.”",
    "Don’t show a real exam platform / brand logos of proctoring vendors.",
    "Don’t advertise Ollama/local AI — it’s not in the beta.",
]:
    bp=d2.add_paragraph(); bp.paragraph_format.space_after=Pt(2); run(bp,"• "+line,size=9.5)
spacer(4)

pagebreak()

# ════════════════════════════════════════════════════════════════════
#  15. WEBSITE STRUCTURE & MESSAGING GUARDRAILS
# ════════════════════════════════════════════════════════════════════
h1("15. Website structure & messaging guardrails")

h2("Recommended page structure")
numbered("the tagline, the 8s hero loop, and one primary CTA (“Start free — 14-day trial” / “Download for Windows”).", bold_prefix="Hero — ")
numbered("the 3-up: Invisible · Screen-aware · Yours (local-first/BYOK). Mirror Group-A capabilities in §2.", bold_prefix="What it does — ")
numbered("the plain-language invisibility explanation from §3 + the honest-limits line. Builds trust.", bold_prefix="How it works — ")
numbered("meetings, learning, interview prep, research (§4) — with screenshots from the mockups.", bold_prefix="Use cases — ")
numbered("provider logos (OpenAI, Anthropic, Google) + “bring your own key.”", bold_prefix="Models / BYOK — ")
numbered("Windows-only, 14-day free trial, sign in with Google, then download. Include the SmartScreen note (§10).", bold_prefix="Get started — ")
numbered("“Interested in a paid plan?” capture (feeds product analytics). Pricing itself is TBD.", bold_prefix="Pricing / interest — ")
numbered("a short, honest privacy summary linking to the full policy (§11).", bold_prefix="Privacy — ")
numbered("Privacy Policy, Beta Terms, support/contact. (Legal copy = counsel-reviewed; see §17.)", bold_prefix="Footer — ")

h2("The proctoring/exam question — read before writing copy")
callout("Strong recommendation",
    [
        "InvisiQ is genuinely invisible to monitoring/proctoring tools — that is a true CAPABILITY. But publicly POSITIONING the product as “beat your proctored exam” creates serious risk: app-store / payment-processor / ad-network rejection, trademark exposure (don’t name Mettl/HackerRank/Respondus/etc.), and brand damage if it reads as a cheating tool.",
        "Recommended public framing: lead with live meetings, interview & assessment PRACTICE, learning, and real-time research. State the monitoring-invisibility as a privacy/capability fact (“stays private even when your screen is shared or monitored”) — not as encouragement to defeat an exam.",
        "This is a recommendation; the final positioning call is the owner’s. Whatever is chosen, keep the language factual and avoid naming specific proctoring vendors or showing their UIs.",
    ],
    tint=TINT_WARN, bar=AMBER, icon="⚠  ")

h2("Claims allow-list (safe, accurate phrasings)")
bullet("“Invisible to screen sharing, screen recording, and screenshots on Windows.”")
bullet("“Stays private even when you’re sharing or your screen is being monitored.”")
bullet("“Bring your own AI key — your screen content never leaves your device.”")
bullet("“Controlled entirely by keyboard shortcuts; always on top, never in your way.”")
h3("Avoid")
bullet("“Undetectable / unhackable / 100% guaranteed invisible.” (Not true vs. physical cameras / kernel DLP.)")
bullet("“Cheat on your exam / bypass your proctor.” (Legal + brand risk.)")
bullet("“Free forever / no AI costs.” (Trial is 14 days; user still pays their provider.)")
bullet("“Runs offline with local models.” (Removed in beta.)")

# ════════════════════════════════════════════════════════════════════
#  16. TECH STACK (credibility / footer)
# ════════════════════════════════════════════════════════════════════
h1("16. Tech stack (for credibility blurbs / FAQ)")
table(
    ["Layer", "Technology"],
    [
        ("Desktop runtime", "Electron 33 (Windows)"),
        ("UI", "React 18 + TypeScript 5 (strict)"),
        ("Styling / motion", "TailwindCSS 3 + Framer Motion 11"),
        ("Build", "electron-vite 5 / electron-builder"),
        ("AI (cloud, BYOK)", "OpenAI, Anthropic, and Google Generative AI SDKs (streaming, lazy-loaded)"),
        ("Vision / OCR", "Direct vision API + Tesseract.js OCR helper"),
        ("Storage", "electron-store + AES-256-GCM (PBKDF2, machine-bound key)"),
        ("Auth / trial / analytics", "Supabase (Google OAuth, server-clock trial, telemetry edge functions)"),
        ("Updates", "electron-updater (auto-update feed)"),
        ("Native helper", "C++ Win32 helper for the suppressing keystroke-capture pipeline"),
    ],
    widths=[2.0, 4.7]
)

# ════════════════════════════════════════════════════════════════════
#  17. OPEN ITEMS / OWNER TO PROVIDE
# ════════════════════════════════════════════════════════════════════
h1("17. Open items the owner must provide before launch")
para("Marked clearly so nothing here is guessed. The website/video should leave placeholders for these.")
table(
    ["Item", "Status / needed for"],
    [
        ("Public download URL / host", ("TBD — owner", "B23B3B"), ),
        ("Support / contact email", ("TBD — owner", "B23B3B")),
        ("Counsel-reviewed Privacy Policy + Beta Terms", ("Required before public launch (§11)", "B23B3B")),
        ("Pricing for post-beta paid plan", ("TBD — owner (site shows “coming soon”)", "B23B3B")),
        ("Geo-availability of the beta (GDPR/DPDP)", ("Owner decision with counsel (§11)", "B23B3B")),
        ("Logo / ghost-mark vector files (SVG)", ("Owner to supply final art (mockups show the intended mark)", "B23B3B")),
        ("Code-signing certificate", ("Deferred until company entity exists; beta is unsigned (§10)", "B23B3B")),
        ("Final positioning on proctoring framing", ("Owner decision (§15)", "B23B3B")),
    ],
    widths=[3.1, 3.6]
)

# ════════════════════════════════════════════════════════════════════
#  18. GLOSSARY
# ════════════════════════════════════════════════════════════════════
h1("18. Quick glossary")
table(
    ["Term", "Meaning"],
    [
        ("WDA_EXCLUDEFROMCAPTURE", "Windows flag that excludes a window from all screen-capture outputs while still showing it on the monitor."),
        ("setContentProtection", "The Electron call that applies that flag — the core of the invisibility."),
        ("BYOK", "Bring Your Own Key — the user supplies their own AI provider API key."),
        ("Stealth typing", "Typing into the overlay from any app without it taking window focus."),
        ("Foreground-window monitoring", "How secure browsers / proctoring tools detect “switching away”; defeated by the non-activatable window."),
        ("Server-clock trial", "The 14-day clock lives on InvisiQ’s server, so it can’t be reset locally."),
    ],
    widths=[2.3, 4.4]
)

spacer(8)
end = para("— End of brief —", align=WD_ALIGN_PARAGRAPH.CENTER, color=SUBTLE, italic=True)

import os
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "InvisiQ-Beta-Launch-Technical-Brief.docx")
doc.save(out)
print("SAVED:", out)
